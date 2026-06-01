"""把上传的简历（PDF/DOCX/TXT）转成纯文本。"""

from __future__ import annotations

import io

from docx import Document
from pypdf import PdfReader


def load_resume(filename: str, blob: bytes) -> str:
    name = (filename or "").lower()
    if name.endswith(".pdf"):
        return _load_pdf(blob)
    if name.endswith(".docx"):
        return _load_docx(blob)
    if name.endswith(".txt") or name.endswith(".md"):
        return blob.decode("utf-8", errors="ignore")
    raise ValueError(f"暂不支持的简历格式: {filename}，请上传 PDF / DOCX / TXT")


def _load_pdf(blob: bytes) -> str:
    reader = PdfReader(io.BytesIO(blob))
    parts: list[str] = []
    for page in reader.pages:
        try:
            parts.append(page.extract_text() or "")
        except Exception:
            continue
    return _normalize("\n".join(parts))


def _load_docx(blob: bytes) -> str:
    doc = Document(io.BytesIO(blob))
    parts = [p.text for p in doc.paragraphs if p.text]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells if cell.text))
    return _normalize("\n".join(parts))


def _normalize(text: str) -> str:
    lines = [ln.strip() for ln in text.splitlines()]
    return "\n".join(ln for ln in lines if ln)
